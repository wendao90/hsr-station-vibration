"""
Convert the paper draft to docx format with embedded figures.
Figure markers in markdown: <!--FIG:filename.png|caption-->
Figures read from: ../analysis/figures/
"""
import os, re, sys, io
try:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
except Exception:
    pass
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, 'manuscript.docx')
FIG_DIR = os.path.join(os.path.dirname(BASE), 'analysis', 'figures')

# ==== Template-aligned typographic constants ====
TITLE_SIZE       = 24.5   # 黑体  bold  center
AUTHOR_SIZE      = 13.0   # 楷体  center
AFFIL_SIZE       = 10.5   # 楷体  center  (Body Text style)
EN_TITLE_SIZE    = 16.0   # Times New Roman  bold  center
EN_AUTHOR_SIZE   = 12.0   # TNR   center
EN_AFFIL_SIZE    = 10.5   # TNR   center
META_SIZE        = 11.0   # CLC + funding + bio lines
SEC_SIZE         = 14.0   # 0, 1, 2, ...    黑体 bold
SUBSEC_SIZE      = 12.0   # 1.1, 2.1, ...   黑体 bold
SUBSUBSEC_SIZE   = 11.0   # 1.1.1          黑体 bold
BODY_SIZE        = 11.0   # 宋体  body
FIG_CAP_SIZE     = 10.5   # 黑体  center
TAB_CAP_SIZE     = 10.5   # 黑体  center
TAB_CELL_SIZE    = 10.0   # 宋体
EQ_SIZE          = 11.0
REF_HEADING_SIZE = 12.0
REF_BODY_SIZE    = 10.5
BODY_INDENT      = Pt(22.0)  # 2 Chinese characters at 11 pt

doc = Document()

for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size_pt=10.5, bold=False, italic=False):
    run.font.name = en_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


INLINE_RE = re.compile(r'(<sub>[^<]+</sub>|<sup>[^<]+</sup>|\*[^*]+\*)')

# ================ OMML equations ================
M_NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'


def _sub(var, sub):
    return f'<m:sSub><m:e><m:r><m:t>{var}</m:t></m:r></m:e><m:sub><m:r><m:t>{sub}</m:t></m:r></m:sub></m:sSub>'


def _subsup(var, sub, sup):
    return (f'<m:sSubSup><m:e><m:r><m:t>{var}</m:t></m:r></m:e>'
            f'<m:sub><m:r><m:t>{sub}</m:t></m:r></m:sub>'
            f'<m:sup><m:r><m:t>{sup}</m:t></m:r></m:sup></m:sSubSup>')


def _space(s=' '):
    return f'<m:r><m:t xml:space="preserve">{s}</m:t></m:r>'


OMML_EQ1 = (
    f'<m:oMath {M_NS}>'
    + _sub('T', 'start') + _space(' = ') + _sub('T', 'audio') + _space(' − ') + _sub('t', 'vib')
    + '</m:oMath>'
)

# a_w = √( W_k² · a_z² + (1.4)² · (W_d² · a_x² + W_d² · a_y²) )
OMML_EQ2 = (
    f'<m:oMath {M_NS}>'
    + _sub('a', 'w') + _space(' = ')
    + '<m:rad><m:radPr><m:degHide m:val="1"/><m:ctrlPr/></m:radPr><m:deg/><m:e>'
    +     _subsup('W', 'k', '2') + _subsup('a', 'z', '2') + _space(' + ')
    +     '<m:sSup><m:e><m:d><m:e><m:r><m:t>1.4</m:t></m:r></m:e></m:d></m:e>'
    +     '<m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>'
    +     '<m:d><m:e>'
    +         _subsup('W', 'd', '2') + _subsup('a', 'x', '2') + _space(' + ')
    +         _subsup('W', 'd', '2') + _subsup('a', 'y', '2')
    +     '</m:e></m:d>'
    + '</m:e></m:rad>'
    + '</m:oMath>'
)


def add_equation(omml_xml, number):
    """Centered OMML equation with right-aligned Chinese-paren number, template style."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Center tab at page midpoint, right tab at right text edge
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(8.5), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run('\t')
    set_run_font(r1, size_pt=EQ_SIZE)
    p._p.append(parse_xml(omml_xml))
    r2 = p.add_run(f'\t（{number}）')
    set_run_font(r2, size_pt=EQ_SIZE, bold=True)
    return p



def add_para(text, align='left', size_pt=BODY_SIZE, cn_font='宋体', bold=False, italic=False,
             first_line_indent=False, parse_inline=False):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if first_line_indent:
        p.paragraph_format.first_line_indent = BODY_INDENT
    if not parse_inline:
        r = p.add_run(text)
        set_run_font(r, cn_font=cn_font, size_pt=size_pt, bold=bold, italic=italic)
        return p
    # Split inline and render sub/sup/italic runs
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            set_run_font(r, cn_font=cn_font, size_pt=size_pt, bold=bold, italic=italic)
        token = m.group(0)
        if token.startswith('<sub>'):
            inner = token[5:-6]
            r = p.add_run(inner)
            set_run_font(r, cn_font=cn_font, size_pt=size_pt, bold=bold, italic=italic)
            r.font.subscript = True
        elif token.startswith('<sup>'):
            inner = token[5:-6]
            r = p.add_run(inner)
            set_run_font(r, cn_font=cn_font, size_pt=size_pt, bold=bold, italic=italic)
            r.font.superscript = True
        else:  # *italic*
            inner = token[1:-1]
            r = p.add_run(inner)
            set_run_font(r, cn_font=cn_font, size_pt=size_pt, bold=bold, italic=True)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        set_run_font(r, cn_font=cn_font, size_pt=size_pt, bold=bold, italic=italic)
    return p


def add_figure(filename, caption, width_cm=14):
    path = os.path.join(FIG_DIR, filename)
    if not os.path.exists(path):
        print(f'[WARN] Missing figure: {path}')
        add_para(f'[图缺失: {filename}]', align='center', size_pt=FIG_CAP_SIZE)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))
    # Caption (template: 黑体 center, not bold)
    add_para(caption, align='center', size_pt=FIG_CAP_SIZE, cn_font='黑体')


def _render_inline_into_paragraph(para, text, size_pt=9, cn_font='宋体', bold=False):
    """Render text into an existing paragraph, parsing <sub>/<sup>/*italic*/**bold** markers."""
    # Strip ** bold markers (treated as plain bold setting)
    text = text.replace('**', '')
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            set_run_font(r, cn_font=cn_font, size_pt=size_pt, bold=bold)
        token = m.group(0)
        if token.startswith('<sub>'):
            r = para.add_run(token[5:-6])
            set_run_font(r, cn_font=cn_font, size_pt=size_pt, bold=bold)
            r.font.subscript = True
        elif token.startswith('<sup>'):
            r = para.add_run(token[5:-6])
            set_run_font(r, cn_font=cn_font, size_pt=size_pt, bold=bold)
            r.font.superscript = True
        else:  # *italic*
            r = para.add_run(token[1:-1])
            set_run_font(r, cn_font=cn_font, size_pt=size_pt, bold=bold, italic=True)
        pos = m.end()
    if pos < len(text):
        r = para.add_run(text[pos:])
        set_run_font(r, cn_font=cn_font, size_pt=size_pt, bold=bold)


def add_markdown_table(rows):
    """Insert a docx table from parsed markdown rows (first row = header)."""
    if not rows or len(rows) < 2:
        return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            tc = t.rows[i].cells[j]
            tc.text = ''
            para = tc.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _render_inline_into_paragraph(para, cell.strip(), size_pt=TAB_CELL_SIZE, cn_font='宋体', bold=(i == 0))


# ================ TITLE BLOCK (template-aligned) ================

# Chinese title: 黑体 24.5 pt, bold, center
add_para('桥建合一高铁站车致振动多场景实测分析',
         align='center', size_pt=TITLE_SIZE, cn_font='黑体', bold=True)

# Chinese authors: 楷体 13 pt, center
add_para('罗鸣璋¹，淡丹辉¹，王文钊¹',
         align='center', size_pt=AUTHOR_SIZE, cn_font='楷体')

# Chinese affiliation: 楷体 smaller, center
add_para('(1. 同济大学，上海 200092)',
         align='center', size_pt=AFFIL_SIZE, cn_font='楷体')

# English title: Times New Roman, bold, center
add_para('Train-Induced Vibration at a Bridge-Building Integrated HSR Station: '
         'A Multi-Scenario Field Measurement',
         align='center', size_pt=EN_TITLE_SIZE, cn_font='Times New Roman', bold=True)
add_para('LUO Mingzhang¹, DAN Danhui¹, WANG Wenzhao¹',
         align='center', size_pt=EN_AUTHOR_SIZE, cn_font='Times New Roman')
add_para('(1. Tongji University, Shanghai 200092, China)',
         align='center', size_pt=EN_AFFIL_SIZE, cn_font='Times New Roman')

# CLC code line (left, 11 pt, 宋体)
add_para('Chinese library classification: U291.6   Document code: A',
         size_pt=META_SIZE, cn_font='Times New Roman')

# Funding and bio (11 pt body)
add_para('基金项目：国家自然科学基金项目（52578234）；上海市自然科学基金（24ZR1468300）；'
         '中铁工程设计咨询集团有限公司项目（KY2024A031）',
         size_pt=META_SIZE)
add_para('作者简介：罗鸣璋（1986—），男，工程师。通讯作者：王文钊（1994—），男。',
         size_pt=META_SIZE)

doc.add_paragraph()

# ================ PARSE MARKDOWN BODY ================
with open(os.path.join(BASE, 'paper_draft.md'), encoding='utf-8') as f:
    md = f.read()

# Extract body: from "## 0 引  言" up to "---" that starts the 附: 建议配图清单
body_start = md.find('## 0 引')
body_end = md.find('## 附：建议配图清单')
if body_end == -1:
    body_end = len(md)
# We want to include摘要 and Abstract BEFORE 0 引言
# Find "## 摘 要" position
abstract_start = md.find('## 摘 要')

# We output: 摘要/Abstract first
# Process each from abstract_start to body_end
section_body = md[abstract_start:body_end] if abstract_start != -1 else md[body_start:body_end]

lines = section_body.split('\n')
table_buffer = []


def flush_table():
    global table_buffer
    if table_buffer:
        add_markdown_table(table_buffer)
        table_buffer = []


FIG_RE = re.compile(r'<!--FIG:([^|]+)\|(.+?)-->')
EQ_RE = re.compile(r'<!--EQ:(\d+)-->')
EQ_MAP = {'1': OMML_EQ1, '2': OMML_EQ2}

i = 0
while i < len(lines):
    line = lines[i].rstrip()
    i += 1

    if not line.strip():
        flush_table()
        continue

    # Equation marker
    m = EQ_RE.search(line)
    if m:
        flush_table()
        num = m.group(1)
        if num in EQ_MAP:
            add_equation(EQ_MAP[num], num)
        continue

    # Figure marker
    m = FIG_RE.search(line)
    if m:
        flush_table()
        fname = m.group(1).strip()
        caption = m.group(2).strip()
        add_figure(fname, caption)
        continue

    if line.startswith('## '):
        flush_table()
        title = line[3:].strip()
        if title in ('摘 要', '摘要:'):
            add_para('摘 要：', size_pt=BODY_SIZE, cn_font='黑体', bold=True)
            continue
        if title == 'Abstract':
            add_para('Abstract:', size_pt=BODY_SIZE, bold=True, cn_font='Times New Roman')
            continue
        if title.startswith('Train-Induced'):
            continue
        # Section heading.
        # Numbered "0/1/2/…": 14 pt 黑体 bold (top-level 正文章节)
        # "参考文献" / "数据与代码可用性" 等非编号小节: 12 pt 黑体 bold (同参考文献级别)
        is_numbered_section = bool(re.match(r'^\d+\s', title))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        set_run_font(r, cn_font='黑体',
                     size_pt=SEC_SIZE if is_numbered_section else REF_HEADING_SIZE,
                     bold=True)
        continue

    if line.startswith('### '):
        flush_table()
        title = line[4:].strip()
        # Subsection "1.1 …": 黑体 12 pt bold
        p = doc.add_paragraph()
        r = p.add_run(title)
        set_run_font(r, cn_font='黑体', size_pt=SUBSEC_SIZE, bold=True)
        continue

    if line.startswith('#### '):
        flush_table()
        title = line[5:].strip()
        # Sub-subsection "1.1.1 …": 黑体 11 pt bold
        p = doc.add_paragraph()
        r = p.add_run(title)
        set_run_font(r, cn_font='黑体', size_pt=SUBSUBSEC_SIZE, bold=True)
        continue

    if line.startswith('# '):
        continue

    # Table row
    if line.startswith('|'):
        # Skip markdown separator rows like |---|---|
        content = line.strip('|').strip()
        if set(content.replace('|', '').strip()) <= {'-', ' ', ':'}:
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        table_buffer.append(cells)
        continue
    else:
        flush_table()

    # Table caption (black body, center, not bold per template)
    if line.startswith('**表') or line.startswith('**Table'):
        caption = line.replace('**', '').strip()
        add_para(caption, align='center', size_pt=TAB_CAP_SIZE, cn_font='黑体')
        continue

    # Bulleted list item (rare in current draft)
    if line.startswith('- ') or re.match(r'^\d+\.\s', line):
        text = re.sub(r'^(\d+\.\s|-\s)', '', line)
        text = text.replace('**', '')
        add_para(('  • ' if line.startswith('-') else '  ') + text,
                 size_pt=BODY_SIZE)
        continue

    if line.strip().startswith('---'):
        doc.add_paragraph()
        continue

    if line.startswith('**作者') or line.startswith('**通讯') or line.startswith('**基金'):
        text = line.replace('**', '')
        add_para(text, size_pt=META_SIZE)
        continue

    # Regular paragraph (11 pt body, first-line indent 2 Chinese chars)
    text = line
    text = text.replace('&nbsp;', '\u00A0')
    text = text.replace('**', '')
    # References in paper_draft use plain [N] lines; they appear as regular paras
    in_refs_body = text.lstrip().startswith('[') and ']' in text[:6]
    if in_refs_body:
        add_para(text, size_pt=REF_BODY_SIZE, parse_inline=True)
    else:
        add_para(text, size_pt=BODY_SIZE, first_line_indent=True, parse_inline=True)


flush_table()

# Save
doc.save(OUT)
print(f'[OK] Manuscript saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes ({os.path.getsize(OUT)/1024:.1f} KB)')
